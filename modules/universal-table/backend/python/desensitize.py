#!/usr/bin/env python3
import argparse
import csv
import json
import os
import random
import re
import sys
from typing import Any, Dict, List


def load_json(value: str, default: Any) -> Any:
    if not value:
        return default
    try:
        return json.loads(value)
    except Exception:
        return default


def mask_phone(text: str) -> str:
    return re.sub(r'(1[3-9]\d)\d{4}(\d{4})', r'\1****\2', text)


def mask_id_card(text: str) -> str:
    return re.sub(r'([1-9]\d{2})\d{11,14}([\dXx]{4})', r'\1********\2', text)


def mask_bank_card(text: str) -> str:
    return re.sub(r'(\d{4})\d{8,11}(\d{4})', r'\1********\2', text)


def mask_email(text: str) -> str:
    return re.sub(r'([A-Za-z0-9._%+-]{2})[A-Za-z0-9._%+-]*(@[A-Za-z0-9.-]+\.[A-Za-z]{2,})', r'\1***\2', text)


def mask_name(text: str) -> str:
    if len(text) <= 1:
        return '*'
    return text[0] + '*' * (len(text) - 1)


def apply_text_mask(text: str) -> str:
    text = mask_phone(text)
    text = mask_id_card(text)
    text = mask_bank_card(text)
    text = mask_email(text)
    return text


def mask_text_lines(text: str, plan_fields: List[Dict[str, Any]]) -> str:
    masked_lines = []
    for line in text.splitlines():
        if '：' in line:
            key, value = line.split('：', 1)
            sensitive_type = explicit_type_for(key.strip(), plan_fields)
            masked_lines.append(f"{key}：{mask_value(value.strip(), sensitive_type) if sensitive_type else apply_text_mask(value.strip())}")
        elif ':' in line:
            key, value = line.split(':', 1)
            sensitive_type = explicit_type_for(key.strip(), plan_fields)
            masked_lines.append(f"{key}:{mask_value(value.strip(), sensitive_type) if sensitive_type else apply_text_mask(value.strip())}")
        else:
            masked_lines.append(apply_text_mask(line))
    return '\n'.join(masked_lines)


def explicit_type_for(header: str, plan_fields: List[Dict[str, Any]]) -> str:
    for field in plan_fields:
        if field.get('field') == header:
            return field.get('sensitiveType', '')
    lowered = header.lower()
    if any(token in lowered for token in ['手机', '电话', 'phone', 'mobile']):
        return 'phone'
    if any(token in lowered for token in ['身份证', '证件']):
        return 'idCard'
    if any(token in lowered for token in ['邮箱', 'email']):
        return 'email'
    if any(token in lowered for token in ['卡号', '银行卡', '账号']):
        return 'bankCard'
    if any(token in lowered for token in ['姓名', 'name', '联系人']):
        return 'name'
    return ''


def mask_value(value: Any, sensitive_type: str) -> str:
    text = '' if value is None else str(value)
    if sensitive_type == 'phone':
        return mask_phone(text)
    if sensitive_type == 'idCard':
        return mask_id_card(text)
    if sensitive_type == 'email':
        return mask_email(text)
    if sensitive_type == 'bankCard':
        return mask_bank_card(text)
    if sensitive_type == 'name':
        return mask_name(text)
    return apply_text_mask(text)


def fake_value(header: str, sample: Any, field_type: str) -> Any:
    sample_text = '' if sample is None else str(sample)
    lowered = header.lower()
    if field_type == 'number' or any(token in lowered for token in ['金额', '数量', '分数', '次数']):
        return random.randint(10, 9999)
    if field_type == 'date' or any(token in lowered for token in ['日期', '时间']):
        return f"2026-{random.randint(1,12):02d}-{random.randint(1,28):02d}"
    if any(token in lowered for token in ['手机', '电话', 'phone']):
        return f"13{random.randint(100000000, 999999999)}"
    if any(token in lowered for token in ['身份证', '证件']):
        return f"11010119900101{random.randint(1000, 9999)}"
    if any(token in lowered for token in ['邮箱', 'email']):
        return f"user{random.randint(100,999)}@example.com"
    if any(token in lowered for token in ['姓名', 'name']):
        return random.choice(['张三', '李四', '王五', '赵六'])
    if sample_text:
        return f"测试{header}"
    return f"示例{header}"


def load_rows_from_csv(file_path: str) -> List[Dict[str, Any]]:
    with open(file_path, 'r', encoding='utf-8-sig', errors='ignore', newline='') as handle:
        return list(csv.DictReader(handle))


def save_rows_to_csv(file_path: str, rows: List[Dict[str, Any]]) -> None:
    headers = list(rows[0].keys()) if rows else []
    with open(file_path, 'w', encoding='utf-8-sig', newline='') as handle:
        writer = csv.DictWriter(handle, fieldnames=headers)
        writer.writeheader()
        writer.writerows(rows)


def load_rows_from_xlsx(file_path: str) -> List[Dict[str, Any]]:
    from openpyxl import load_workbook

    workbook = load_workbook(file_path)
    sheet = workbook[workbook.sheetnames[0]]
    data = list(sheet.iter_rows(values_only=True))
    workbook.close()
    if not data:
        return []
    headers = [str(item).strip() if item is not None else '' for item in data[0]]
    rows = []
    for row in data[1:]:
        item = {}
        for index, header in enumerate(headers):
            if header:
                item[header] = row[index] if index < len(row) else ''
        if item:
            rows.append(item)
    return rows


def save_rows_to_xlsx(file_path: str, rows: List[Dict[str, Any]]) -> None:
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    headers = list(rows[0].keys()) if rows else []
    if headers:
        sheet.append(headers)
    for row in rows:
        sheet.append([row.get(header, '') for header in headers])
    workbook.save(file_path)


def load_docx_parts(file_path: str):
    from docx import Document

    doc = Document(file_path)
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        tables.append(rows)
    return paragraphs, tables


def save_docx_parts(file_path: str, paragraphs: List[str], tables: List[List[List[str]]]) -> None:
    from docx import Document

    doc = Document()
    for paragraph in paragraphs:
        if paragraph:
            doc.add_paragraph(paragraph)
    for table_rows in tables:
        if not table_rows:
            continue
        table = doc.add_table(rows=0, cols=len(table_rows[0]))
        for row_values in table_rows:
            row = table.add_row().cells
            for index, value in enumerate(row_values):
                row[index].text = '' if value is None else str(value)
    doc.save(file_path)


def read_pdf_text(file_path: str) -> str:
    import pdfplumber

    parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ''
            if text.strip():
                parts.append(text)
    return '\n'.join(parts)


def save_pdf_text(file_path: str, text: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas

    font_path = os.path.join(os.environ.get('WINDIR', 'C:\\Windows'), 'Fonts', 'msyh.ttc')
    if os.path.exists(font_path):
        pdfmetrics.registerFont(TTFont('MSYH', font_path))
        font_name = 'MSYH'
    else:
        font_name = 'Helvetica'

    c = canvas.Canvas(file_path, pagesize=A4)
    c.setFont(font_name, 10)
    x, y = 40, 800
    for line in text.splitlines():
        c.drawString(x, y, line[:90])
        y -= 14
        if y < 40:
            c.showPage()
            c.setFont(font_name, 10)
            y = 800
    c.save()


def mask_rows(rows: List[Dict[str, Any]], plan_fields: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    result = []
    for row in rows:
        item = {}
        for header, value in row.items():
            sensitive_type = explicit_type_for(header, plan_fields)
            item[header] = mask_value(value, sensitive_type) if sensitive_type else apply_text_mask('' if value is None else str(value))
        result.append(item)
    return result


def build_test_rows(rows: List[Dict[str, Any]], schema: List[Dict[str, Any]], count: int) -> List[Dict[str, Any]]:
    headers = [field.get('title') for field in schema if field.get('title')] or (list(rows[0].keys()) if rows else [])
    output = []
    for index in range(count):
        record = {}
        source = rows[index % len(rows)] if rows else {}
        for field in schema:
            title = field.get('title')
            if not title:
                continue
            record[title] = fake_value(title, source.get(title), field.get('type', 'string'))
        if not schema:
            for header in headers:
                record[header] = fake_value(header, source.get(header), 'string')
        output.append(record)
    return output


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument('--action', required=True, choices=['mask', 'testdata'])
    parser.add_argument('--input', required=True)
    parser.add_argument('--output', required=True)
    parser.add_argument('--mode', default='rule')
    parser.add_argument('--plan-json', default='{}')
    parser.add_argument('--schema-json', default='[]')
    parser.add_argument('--count', type=int, default=10)
    args = parser.parse_args()

    ext = os.path.splitext(args.input)[1].lower()
    plan = load_json(args.plan_json, {})
    plan_fields = plan.get('fields', []) if isinstance(plan, dict) else []
    schema = load_json(args.schema_json, [])

    if ext == '.csv':
        rows = load_rows_from_csv(args.input)
        output_rows = mask_rows(rows, plan_fields) if args.action == 'mask' else build_test_rows(rows, schema, args.count)
        save_rows_to_csv(args.output, output_rows)
    elif ext in ('.xlsx', '.xlsm'):
        rows = load_rows_from_xlsx(args.input)
        output_rows = mask_rows(rows, plan_fields) if args.action == 'mask' else build_test_rows(rows, schema, args.count)
        save_rows_to_xlsx(args.output, output_rows)
    elif ext == '.docx':
        paragraphs, tables = load_docx_parts(args.input)
        if args.action == 'mask':
            masked_paragraphs = [mask_text_lines(text, plan_fields) for text in paragraphs]
            masked_tables = [[[(apply_text_mask(cell)) for cell in row] for row in table] for table in tables]
            save_docx_parts(args.output, masked_paragraphs, masked_tables)
        else:
            rows = []
            if tables:
                headers = tables[0][0]
                for table in tables:
                    for row in table[1:]:
                        rows.append({headers[i]: row[i] if i < len(row) else '' for i in range(len(headers))})
            output_rows = build_test_rows(rows, schema, args.count)
            table_rows = [list(output_rows[0].keys())] + [list(row.values()) for row in output_rows]
            save_docx_parts(args.output, [], [table_rows])
    elif ext == '.pdf':
        text = read_pdf_text(args.input)
        if args.action == 'mask':
            save_pdf_text(args.output, mask_text_lines(text, plan_fields))
        else:
            rows = build_test_rows([], schema, max(1, args.count))
            content = '\n'.join([' | '.join(row.keys()), *[' | '.join(str(value) for value in row.values()) for row in rows]])
            save_pdf_text(args.output, content)
    else:
        with open(args.input, 'r', encoding='utf-8', errors='ignore') as handle:
            text = handle.read()
        if args.action == 'mask':
            output_text = mask_text_lines(text, plan_fields)
        else:
            rows = build_test_rows([], schema, max(1, args.count))
            output_text = '\n'.join(['：'.join([key, str(value)]) for row in rows for key, value in row.items()])
        with open(args.output, 'w', encoding='utf-8') as handle:
            handle.write(output_text)

    print(json.dumps({
        'success': True,
        'outputPath': args.output,
        'outputFileName': os.path.basename(args.output),
    }, ensure_ascii=False))


if __name__ == '__main__':
    main()
